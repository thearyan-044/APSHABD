from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = r"C:\Users\LENOVO\OneDrive\Aryan\side\PDS\PDS Brand Doc.docx"
NAVY=RGBColor(11,37,69); BLUE=RGBColor(46,116,181); GRAY=RGBColor(90,98,108); WHITE=RGBColor(255,255,255)

doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor(32,36,42)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,before,after,color in [('Title',30,0,8,NAVY),('Subtitle',13,0,18,GRAY),('Heading 1',16,18,10,BLUE),('Heading 2',13,14,7,BLUE),('Heading 3',12,10,5,RGBColor(31,77,120))]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=color
    s.font.bold=name!='Subtitle'; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

header=sec.header.paragraphs[0]; header.text='PIN DROP SILENCE  /  BRAND DOC'; header.style=styles['Caption']
header.runs[0].font.color.rgb=GRAY
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def add_bullet(text,level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.add_run(text); return p
def add_num(text):
    p=doc.add_paragraph(style='List Number'); p.add_run(text); return p
def add_quote(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.3); p.paragraph_format.right_indent=Inches(.2)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text); r.bold=True; r.italic=True; r.font.color.rgb=NAVY
    return p
def add_status(value):
    p=doc.add_paragraph(); r=p.add_run('STATUS  '); r.bold=True; r.font.color.rgb=BLUE; p.add_run(value)
def heading(text,level=1): doc.add_heading(text,level=level)
def para(text,boldlead=None):
    p=doc.add_paragraph()
    if boldlead and text.startswith(boldlead):
        p.add_run(boldlead).bold=True; p.add_run(text[len(boldlead):])
    else: p.add_run(text)
    return p
def bullets(items):
    for x in items: add_bullet(x)

# Editorial cover
for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PDS'); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE
p=doc.add_paragraph('PDS Brand Doc',style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph('Pin Drop Silence — Living Brand Strategy',style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph('Seven Cities. One Silence.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True; p.runs[0].font.size=Pt(16); p.runs[0].font.color.rgb=NAVY
for _ in range(7): doc.add_paragraph()
p=doc.add_paragraph('Local Word copy of the Notion source • Retrieved 19 August 2026'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.color.rgb=GRAY
doc.add_page_break()

heading('Document status')
t=doc.add_table(rows=0,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
for a,b in [('Current','Posting-ready'),('Completed','6 of 18 sections'),('Status','Fast-track in progress'),('Last website review','13 August 2026'),('Notion status','In progress')]:
    c=t.add_row().cells; c[0].width=Inches(1.875); c[1].width=Inches(4.625); c[0].text=a; c[1].text=b; c[0].paragraphs[0].runs[0].bold=True
for row in t.rows:
    for c in row.cells: c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
para('Living document: This is the single working source for Pin Drop Silence brand decisions. Website evidence, working hypotheses and approved decisions remain clearly separated.')

heading('Current Website Evidence')
bullets([
'Master idea: “Seven Cities. One Silence.”',
'Category: A streetwear house for Indian streets that “speak without shouting.”',
'Geographic world: Chennai, Mumbai, Delhi, Pune, Bangalore, Kolkata and Hyderabad.',
'Product concept: Neighbourhood-led graphic tees and varsity pieces using Indian scripts and Western streetwear forms.',
'Emotional material: The roads people grew up on, personal memory, routines, relationships, work and belonging.',
'Community signal: Visitors are invited to suggest locations for future drops.',
'Current scale signal: 50 designs described as in development—30 tees and 20 neighbourhood varsity pieces.'
])
add_quote('Working interpretation—not yet approved: PDS is a hyperlocal Indian streetwear brand built around neighbourhood identity and understated cultural expression, rather than a generic graphic T-shirt or tourist-merchandise label.')

heading('1. Brand Problem and Opportunity'); add_status('Complete')
heading('Website evidence',2); bullets(['PDS is “made for Indian streets” and says those streets “speak without shouting.”','Every piece is intended to carry a neighbourhood rather than borrowed Western locations.','The site combines modern Western streetwear with Indian streets, scripts and lived memories.'])
heading('Approved problem',2)
para('Indian streetwear often borrows heavily from Western hip-hop and rap culture. PDS does not see that influence as inherently negative—the founder personally enjoys and participates in that culture. The gap is that Indian identity, places and emotional roots are frequently absent from the resulting clothes.')
para('Global mobility, education and economic opportunity can also create cultural distance for Indians living both inside and outside India. PDS sees an opportunity to make connection to Indian roots feel contemporary and desirable, without reducing Indian identity to festival wear, costume or conventional patriotic merchandise.')
heading('Approved opportunity',2); add_quote('PDS brings Western streetwear language and Indian cultural identity together as equals, allowing both to run in parallel.')
para('The aim is not to reject Western influence or tell people how Indian they should be. It is to create clothing through which people can remember, feel and wear their roots while retaining the silhouettes and cultural codes they already love.')
heading('Strategic guardrails',2); bullets(['Fusion, not rejection: Western streetwear is a creative language, not the enemy.','Roots, not nationalism: Foster cultural connection without becoming political, exclusionary or preachy.','Everyday identity, not occasion wear: Indian feeling should live in daily streetwear.','Parallel cultures: Neither Western nor Indian culture should feel pasted onto or subordinate to the other.','Emotion beyond graphics: Roots must shape stories, neighbourhood choices, language and community.','India and diaspora: The emotional territory can serve Indians at home and abroad.'])
heading('Founder origin — approved',2)
para('When the founder was in fourth grade, he was walking through a Mumbai mall with his father when they noticed someone wearing a good-looking oversized Brooklyn T-shirt. His father asked why there was no Indian-origin equivalent—something with Thane, Powai, Bandra or another Indian neighbourhood in place of Brooklyn.')
para('That observation stayed with him: Indian places could carry the same design confidence and cultural value as globally repeated locations. He later reached a point where he had both the courage and time to build it.')
add_quote('Origin in one line: PDS began with a childhood question: why do we wear Brooklyn when our own neighbourhoods deserve to be worn?')
heading('Strategic meaning',2); bullets(['PDS was born from an observation about local representation, not personal frustration.','The original creative substitution is Brooklyn → Thane, Powai, Bandra and other Indian neighbourhoods.','The garment must look desirable first; Indian identity should strengthen the design rather than feel like novelty merchandise.','The idea stayed with the founder from childhood until he was ready to execute it.','The story supports: “Not Brooklyn. Not NY. Your area—in your script, on your back.”'])
heading('Desired wearer experience — approved',2)
para('A person wearing PDS should feel part of home and connected to their roots while enjoying Western culture, music and streetwear. The garment should reactivate the small, specific memories that made a neighbourhood theirs.')
add_quote('Emotional promise: Wherever you go and whoever you become, the place that made you still belongs to you—and you still belong to it.')
heading('What the wearer communicates',2); bullets(['I remember where I came from.','My neighbourhood helped shape who I became.','I can participate in global culture without erasing my roots.','My local story deserves the same design status as an internationally recognised place.','I am proud of the distance I have travelled without being ashamed of my starting point.'])
heading('Unity principle',2); para('PDS should emphasise the shared experience of place over divisions created by religion, politics or social identity. Express this through human stories and shared local culture—not party politics, propaganda or commentary on specific governments.')
heading('Creative implication',2); para('A neighbourhood design cannot rely only on its name or script. It should contain recognisable details, symbols, incidents or emotional cues that allow someone from that place to say, “This knows my area.”')
heading('Emotional territory hierarchy — approved',2)
for x in ['City pride — primary promise and broadest source of identification.','Neighbourhood belonging — specific, credible expression of city pride.','Personal memory — emotional evidence that makes belonging meaningful.','Quiet confidence — the manner in which the first three are expressed.']: add_num(x)
add_quote('Brand hierarchy: City pride → neighbourhood belonging → personal memory, expressed with quiet confidence.')
heading('Non-negotiable brand boundary — approved direction',2)
para('PDS should be exclusive rather than universally available. Access should favour people who intentionally sign up, follow drops and want to belong. It must not become a default status product bought only for a logo or social signal.')
heading('Strategic translation',2); bullets(['Earned access: Signing up, following a drop and paying attention should matter.','Scarcity: Product availability should remain controlled.','Cultural credibility: Wearing PDS should signal genuine affinity.','Substance over flex: Location story and garment quality matter more than logo visibility.','Community before reach: Not every possible sale is automatically good for the brand.','Customer-facing language: Communicate belonging without insulting or policing customers.'])
heading('Dual exclusivity model — approved',2); bullets(['Premium pricing: Signals material quality, design value and controlled availability.','Cultural membership: Belonging also depends on attention, early sign-up, participation and authentic connection.'])
add_quote('Exclusivity principle: PDS is premium enough to be considered and culturally specific enough to be understood—not merely purchased.')

heading('2. Brand Positioning'); add_status('Complete')
heading('Public category — approved',2); add_quote('PDS is a premium city-led streetwear house.')
heading('Meaning of the category',2); bullets(['Premium: Quality, pricing, scarcity and experience sit above mass fashion.','City-led: Cities are primary; neighbourhoods provide product specificity.','Streetwear: Contemporary silhouettes and cultural codes—not souvenir or occasion wear.','House: An enduring design world and multiple product categories.'])
heading('Usage rule',2); para('Use “premium city-led streetwear house” publicly. Short creative copy may use “city-led streetwear.” Do not position PDS as city merchandise, a souvenir brand, a patriotic label, a generic premium T-shirt brand, or Western streetwear with Indian graphics added.')
heading('Core age range — approved',2); bullets(['Primary wearer: 15–25.','Flexible edge: approximately 26.','Secondary wearer: Anyone older who genuinely connects and is comfortable in the silhouettes.','Principle: Youth-focused, not age-exclusive.'])
heading('Geographic priority — approved',2); bullets(['Phase 1: Chennai, Mumbai, Delhi, Pune, Bangalore, Kolkata and Hyderabad.','Phase 2: Selected Tier 2 cities after the opening model is proven.','Phase 3: Tier 3 cities later, based on credible storytelling, community interest and premium demand.','India comes before international expansion; diaspora is a future opportunity.'])
heading('Life-stage approach — approved',2); para('Treat school students 15–18, college students 18–22 and early professionals 22–25 as one core youth audience until launch data supports segmentation.')
heading('Cultural and behavioural profile — approved',2); bullets(['Interpretive curiosity.','Lived local connection.','Desire to wear that connection.'])
heading('Cultural world and occasions',2); bullets(['Western and Indian music, concerts and live events.','Parties, house parties and urban nightlife.','Art, visual design and creative internet culture.','Local food, routes, landmarks, slang and neighbourhood rituals.','Daily casual wear, campus life and informal cultural events.'])
add_quote('Wearer shorthand: An everyday urban Indian with interpretive taste, a lived relationship with their city and a desire to carry that place into contemporary social life.')
heading('Decisive difference — working hypothesis',2); add_quote('PDS turns a person’s city identity into a visually intriguing garment that attracts attention, raises questions and makes people want to discover the story behind it.')
heading('Attention hierarchy — approved',2)
for x in ['Overall artwork','Product scarcity','Recognisable neighbourhood detail']: add_num(x)
heading('Positioning statement — approved',2); add_quote('PDS is a premium city-led streetwear house for Indians aged 15–25. It creates limited streetwear based on Indian cities and neighbourhoods. The artwork is designed to attract attention, start conversations, and connect wearers to where they come from.')

heading('3. Target Customer'); add_status('Complete')
heading('Approved customer',2); bullets(['The wearer buys and pays for the product.','Living in the featured place is not required; appreciation is enough.','Customers can submit location suggestions through a form.','Existing and first-time premium-streetwear buyers are included.'])
heading('Purchase barriers — approved',2); bullets(['Product scarcity or unavailability.','Designs may not appeal to everyone or may offend some people.','Some may dislike the brand idea, find it uncool or cringe.','PDS has not created hype yet.','No single main barrier is confirmed.'])
heading('Target customer summary — approved',2); bullets(['Primary age: 15–25.','Initial geography: seven launch cities in India.','Appreciates art and interpretation.','Wears PDS for daily wear, concerts, parties and nightlife.'])

heading('4. Consumer Insight'); add_status('Complete')
heading('Purchase triggers — approved',2)
for x in ['Association with the featured area or city.','FOMO created by product scarcity.']: add_num(x)
heading('Consumer insight statement — approved',2); add_quote('I like contemporary streetwear, but I rarely see Indian places I care about represented in a way I want to wear. When PDS creates a limited design around one of those places, I want to get it before it is unavailable.')

heading('5. Brand Promise and Differentiation'); add_status('Complete with deferred items')
bullets(['Core brand promise: Deferred.','Differentiation pillars: Product quality and community.','Community difference: Collaboration and relatability to cities, neighbourhoods and experiences.','Collaboration spaces: Instagram channels and Discord servers.','Approved community designs earn contributors a 5% commission; operating details are deferred.','Product-quality specifics: Deferred.','Public priority: 1) Community, 2) Product quality.'])
heading('Differentiation statement — approved',2); add_quote('PDS builds streetwear with its community, not only for it. Community members can contribute designs, and approved contributors earn commission. Product quality supports the result.')

heading('6. Brand Identity and Voice'); add_status('Complete with deferred item')
heading('Brand personality — approved',2); bullets(['Vibrant','Concert','Energetic','Adrenaline','Trendy'])
heading('Energy split and written voice',2); bullets(['Visual identity: vibrant and energetic.','Written voice: quieter; sarcastic, monotone and rebellious.','Practical definition of “monotone”: Deferred.'])
heading('Public language — approved',2); bullets(['Profanity, vulgar language and slang are required.','Every public message must contain at least one profanity; longer copy uses it only in selected lines.','Never use slurs targeting caste, religion, race, gender, sexuality or disability.','Sarcasm may target situations, culture and PDS itself—never customers.'])
heading('Rebellious and political commentary — approved',2); bullets(['Challenge false freedom, misleading policies, propaganda and blind conformity.','Keep criticism general; do not name governments, political parties or institutions.','Current-event commentary may not name actors.','Use only verified facts; no speculation or unverified claims.','For sensitive topics, keep sources in the internal approval record.','For non-sensitive commentary, reference sources in the post and provide an accessible link.'])
heading('Customer-service and transactional voice — approved',2); bullets(['Use less sarcasm; jokes may be used; remain casual, appealing and non-offensive.','Do not use profanity in service or transactional messages.','No jokes for missing orders, failed payments, refund delays or angry complaints.','Use professional language in high-stakes situations.','When PDS is at fault, apologise directly, name the exact mistake and state the fix.'])

heading('7. Product Strategy'); add_status('In progress')
heading('Launch product category — revised and approved',2); bullets(['Launch with oversized, unisex tees only.','Varsity pieces will be announced later.','Multiple designs and design-specific colour variants.','Each version is assigned to a city or location.','Sizes are not yet defined.'])
heading('Initial design allocation — approved',2); bullets(['15 featured locations per city.','One tee design per location.','15 tee designs per city.'])
heading('Colour, material and decoration',2); bullets(['Four to five colour options per design; selected designs may have six.','Cotton only; it must be high quality.','Fabric weight: Deferred until the full brand document is completed.','Primary decoration method: Screen printing.','Mixed decoration methods may be used; additional methods are undefined.'])
heading('Design variants and print placement — approved',2); bullets(['Minimalistic','Balanced','Front-heavy','Back-heavy','Each location receives one variant at launch.'])
heading('Open question 12',2); add_quote('How should PDS decide which of the four variants a location receives?')

heading('8–18. Pending sections')
for x in ['8. City and Neighbourhood Drop Architecture','9. Pricing and Value Perception','10. D2C Buying Experience','11. Go-to-Market Strategy','12. Content and Community','13. Retention and Loyalty','14. Operations and Customer Trust','15. Unit Economics and Measurement','16. Competitive Landscape','17. Cultural Responsibility and Brand Protection','18. Launch Plan and Milestones']: add_bullet(x+' — Pending.')

heading('Working Rules')
heading('Website',2); bullets(['Re-check the PDS website before drafting or revising a section.','Treat website language as reference evidence, not an approved decision.'])
heading('Writing format',2); bullets(['Write point to point.','Prefer short bullets over paragraphs.','Use plain, literal language.','Keep only information needed for future decisions or content generation.','Avoid jazzy, decorative or overly descriptive phrasing.','Do not add strategic interpretation unless the founder explicitly requests it.'])
heading('Accuracy and updates',2); bullets(['Do not assume missing details or expand founder answers.','Ask a direct follow-up if an answer is unclear, incomplete or contradictory.','Keep approved decisions separate from evidence and open questions.','Label unavoidable inference as Unapproved inference.','Update progress and decision log after each decision.','Preserve previous decisions and show revisions transparently.'])
heading('Fast-track mode — approved',2); bullets(['Prioritise decisions required to publish content immediately.','Ask remaining posting questions in one batch.','Defer non-essential product, pricing, operations and measurement details.','Return to deferred sections after posting has started.'])
heading('Pre-launch posting direction — approved',2); bullets(['Keep initial posts city-neutral and intentionally vague.','Do not reveal a specific city yet.','Primary goals: awareness and community sign-ups.','Selling tees is secondary.','Main conversion: join the waitlist.','Main trigger: FOMO created by scarcity.','Product photographs are ready.','Target posting start: 14 or 15 August 2026.','Cadence: 7–9 posts per week; one daily plus up to two extras.'])

doc.add_paragraph()
p=doc.add_paragraph('Source note: Local Word copy created from the Notion page “PDS Brand Doc,” retrieved 19 August 2026. The Notion page remains the living source of truth.'); p.runs[0].italic=True; p.runs[0].font.color.rgb=GRAY

doc.core_properties.title='PDS Brand Doc'
doc.core_properties.subject='Pin Drop Silence living brand strategy'
doc.core_properties.author='Pin Drop Silence'
doc.save(OUT)
print(OUT)
